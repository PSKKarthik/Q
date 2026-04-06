from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
INPUT_FILE = ROOT / "QGX_DISSERTATION.txt"
OUTPUT_FILE = ROOT / "QGX_DISSERTATION.docx"


@dataclass
class FormatState:
    size_pt: int = 12
    bold: bool = False
    italic: bool = False
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY
    first_line_indent_cm: float = 0.9
    font_name: str = "Times New Roman"


def add_page_number(paragraph, alignment: WD_ALIGN_PARAGRAPH) -> None:
    paragraph.alignment = alignment
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def set_page_number_format(section, fmt: str, start: int | None = None) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    sect_pr.append(pg_num_type)


def configure_section(section, footer_alignment: WD_ALIGN_PARAGRAPH, page_fmt: str, start: int | None = None) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.81)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    set_page_number_format(section, page_fmt, start=start)
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        paragraph = footer.paragraphs[0]
        paragraph.clear()
    else:
        paragraph = footer.add_paragraph()
    add_page_number(paragraph, footer_alignment)


def extract_font_state(line: str, state: FormatState) -> None:
    size_match = re.search(r"(\d+)pt", line, flags=re.IGNORECASE)
    if size_match:
        state.size_pt = int(size_match.group(1))
    if re.search(r"\bBold\b", line, flags=re.IGNORECASE):
        state.bold = True
    if re.search(r"\bNormal\b", line, flags=re.IGNORECASE):
        state.bold = False
    if re.search(r"\bItalic\b", line, flags=re.IGNORECASE):
        state.italic = True
    if re.search(r"\bJustified?\b", line, flags=re.IGNORECASE):
        state.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    indent_match = re.search(r"First-line indent\s*([0-9.]+)cm", line, flags=re.IGNORECASE)
    if indent_match:
        state.first_line_indent_cm = float(indent_match.group(1))


def extract_alignment(line: str, state: FormatState) -> None:
    upper = line.upper()
    if "RIGHT ALIGNMENT" in upper or "BOTTOM RIGHT" in upper:
        state.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif "BOTTOM CENTER" in upper or "CENTER ALIGNMENT" in upper or "CENTER ALIGNED" in upper or upper.strip() == "[CENTER]":
        state.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif "TOP LEFT" in upper or "LEFT ALIGNMENT" in upper or "LEFT-CENTER ALIGNMENT" in upper or "LEFT-FLUSH" in upper:
        state.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(document: Document, text: str, state: FormatState, *, monospace: bool = False, italic: bool | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = state.alignment
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    if state.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and not monospace:
        paragraph.paragraph_format.first_line_indent = Cm(state.first_line_indent_cm)
    else:
        paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(text)
    run.font.name = "Courier New" if monospace else state.font_name
    run.font.size = Pt(11 if monospace else state.size_pt)
    run.bold = state.bold and not monospace
    run.italic = state.italic if italic is None else italic


def main() -> None:
    if not INPUT_FILE.exists():
        raise FileNotFoundError(f"Input file not found: {INPUT_FILE}")

    lines = INPUT_FILE.read_text(encoding="utf-8").splitlines()

    document = Document()
    configure_section(
        document.sections[0],
        footer_alignment=WD_ALIGN_PARAGRAPH.CENTER,
        page_fmt="upperRoman",
        start=1,
    )

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    state = FormatState()
    first_page_marker_seen = False
    code_block_mode = False
    in_main_body = False

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if not stripped:
            continue

        if stripped == "CHAPTER PAGES BEGIN HERE – Arabic page numbering, bottom-right corner":
            document.add_section(WD_SECTION.NEW_PAGE)
            configure_section(
                document.sections[-1],
                footer_alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                page_fmt="decimal",
                start=1,
            )
            in_main_body = True
            continue

        if stripped.startswith("PAGE "):
            if first_page_marker_seen:
                document.add_page_break()
            first_page_marker_seen = True
            code_block_mode = False
            continue

        if stripped.startswith("FORMAT :") or stripped.startswith("Paper  :") or stripped.startswith("Margins:"):
            continue

        if set(stripped) == {"="} or set(stripped) == {"-"}:
            continue

        if stripped.startswith("[") and stripped.endswith("]"):
            code_block_mode = False if stripped.startswith("[BLANK LINE") else code_block_mode

            if stripped.startswith("[FONT:"):
                extract_font_state(stripped, state)
                continue
            if "ALIGNMENT" in stripped.upper() or stripped.upper() in {"[CENTER]", "[TOP LEFT]", "[BOTTOM RIGHT]", "[BOTTOM CENTER]"}:
                extract_alignment(stripped, state)
                continue
            if stripped.startswith("[BLANK LINE"):
                count_match = re.search(r"x(\d+)", stripped, flags=re.IGNORECASE)
                count = int(count_match.group(1)) if count_match else 1
                for _ in range(count):
                    document.add_paragraph("")
                continue
            if stripped.startswith("[PAGE NUMBER:") or stripped.startswith("[NO PAGE NUMBER"):
                continue
            if stripped.startswith("[FIGURE PLACEHOLDER") or stripped.startswith("[TABLE STRUCTURE") or stripped.startswith("[TABLE CAPTION"):
                continue
            if stripped.startswith("[CODE BLOCK"):
                code_block_mode = True
                continue
            if stripped.startswith("[PLACEHOLDER:") or stripped.startswith("[INSERT "):
                placeholder_text = stripped[1:-1]
                add_paragraph(document, placeholder_text, state, italic=True)
                continue
            if stripped.startswith("[TWO-COLUMN LAYOUT"):
                state.alignment = WD_ALIGN_PARAGRAPH.LEFT
                continue
            continue

        if stripped == "END OF DOCUMENT":
            add_paragraph(document, stripped, state)
            continue

        if code_block_mode or stripped.startswith("+") or stripped.startswith("|") or line.startswith("  "):
            add_paragraph(document, line, state, monospace=True)
            continue

        if stripped.endswith(":") and stripped.upper() == stripped:
            previous_size = state.size_pt
            previous_bold = state.bold
            previous_alignment = state.alignment
            state.size_pt = 14
            state.bold = True
            state.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_paragraph(document, stripped, state)
            state.size_pt = previous_size
            state.bold = previous_bold
            state.alignment = previous_alignment
            continue

        if stripped.startswith("FINAL MANDATORY FORMATTING RULES"):
            previous_size = state.size_pt
            previous_bold = state.bold
            previous_alignment = state.alignment
            state.size_pt = 14
            state.bold = True
            state.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_paragraph(document, stripped, state)
            state.size_pt = previous_size
            state.bold = previous_bold
            state.alignment = previous_alignment
            continue

        if stripped.startswith("WORD COUNT SUMMARY") or stripped.startswith("FORMATTING NOTES FOR MS-WORD"):
            previous_size = state.size_pt
            previous_bold = state.bold
            previous_alignment = state.alignment
            state.size_pt = 12
            state.bold = True
            state.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_paragraph(document, stripped, state)
            state.size_pt = previous_size
            state.bold = previous_bold
            state.alignment = previous_alignment
            continue

        if in_main_body and state.size_pt == 12 and not state.bold and state.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            state.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        add_paragraph(document, stripped, state)

    document.save(OUTPUT_FILE)
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()